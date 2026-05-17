"""
Автозаполнение паспорта обследования артезианской скважины (DOCX) и экспорт в PDF.
Шаблон: static/templates/well_passport_template.docx
"""
from __future__ import annotations

import json
import logging
import os
import re
import shutil
import subprocess
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

from docx import Document

_BASE_DIR = Path(__file__).resolve().parent
TEMPLATE_PATH = _BASE_DIR / 'static' / 'templates' / 'well_passport_template.docx'

DEFAULT_EXECUTOR = os.environ.get(
    'WELL_PASSPORT_EXECUTOR', 'Индивидуальный предприниматель Емельянов А.Н.'
)


class PdfConversionError(RuntimeError):
    pass


def _num(v: Any) -> Optional[float]:
    if v is None or v == '':
        return None
    try:
        x = float(str(v).replace(',', '.'))
        return x if x == x else None
    except (TypeError, ValueError):
        return None


def fmt_num_comma(val: Optional[float], decimals: int = 1) -> str:
    if val is None:
        return '—'
    s = f'{val:.{decimals}f}'.rstrip('0').rstrip('.')
    return s.replace('.', ',')


def format_duration_ru(seconds: Optional[float]) -> str:
    if seconds is None or seconds <= 0:
        return '1 часа'
    sec = float(seconds)
    if sec >= 3600:
        h = sec / 3600
        hi = int(round(h))
        if hi == 1:
            return '1 часа'
        if 2 <= hi <= 4:
            return f'{hi} часа'
        return f'{hi} часов'
    minutes = max(1, int(round(sec / 60)))
    if minutes == 1:
        return '1 минуты'
    if 2 <= minutes <= 4:
        return f'{minutes} минуты'
    return f'{minutes} минут'


def _debit_from_measure(inputs: dict) -> Optional[float]:
    liters = _num(inputs.get('measure_liters'))
    seconds = _num(inputs.get('measure_seconds'))
    if liters is None or seconds is None or seconds <= 0:
        return None
    return (liters / 1000) / (seconds / 3600)


def _debit_from_pump(inputs: dict) -> Optional[float]:
    stat = _num(inputs.get('static_m'))
    dyn = _num(inputs.get('dynamic_m'))
    calc = _num(inputs.get('calc_depth_m'))
    pump = _num(inputs.get('pump_m3h'))
    if None in (stat, dyn, calc, pump) or abs(dyn - stat) < 1e-9:
        return None
    return pump / (dyn - stat) * (calc - stat)


def _pick_str(overrides: dict, inputs: dict, key: str, default: str = '') -> str:
    for src in (overrides, inputs):
        v = src.get(key)
        if v is not None and str(v).strip():
            return str(v).strip()
    return default


def _field_line(label: str, value: str, min_pad: int = 44) -> str:
    v = (value or '').strip()
    if v:
        return f'{label}{v}'
    pad = max(12, min_pad - len(label))
    return f'{label}{"_" * pad}'


def _classify_address_chunk(ch: str) -> Optional[str]:
    s = (ch or '').strip()
    if not s:
        return None
    low = s.lower()
    if re.search(r'область|обл\.\s*$', low) or re.match(r'^минск(ая|ий)\s+обл', s, re.I):
        return 'region'
    if re.search(r'район|р-н|р\.\s*н\.?', low):
        return 'district'
    if re.match(r'^ул\.|^\s*улица\b|^пр\.|^\s*проспект\b|^пер\.|^\s*переулок\b|^б-р\b|^\s*бульвар\b', s, re.I):
        return 'street'
    if re.match(r'^д\.\s*[A-Za-zА-Яа-яЁё]', s) and not re.match(r'^д\.\s*\d', s):
        return 'settlement'
    if re.match(r'^аг\.|^п\.|^пос\.|^г\.|^с\.|^дер\.|^\s*деревня\b|^\s*посёлок\b|^\s*поселок\b', s, re.I):
        return 'settlement'
    if re.match(r'^\d+[a-zA-Zа-яА-ЯЁё/-]*$', s) and len(s) <= 10:
        return 'house'
    return None


def _split_street_house(ch: str) -> tuple[str, str]:
    m = re.match(r'^(.+?)\s+(\d+[a-zA-Zа-яА-ЯЁё/-]*)$', ch)
    if m and re.search(r'ул\.|улица|пр\.|пер\.|бульвар|б-р', m.group(1), re.I):
        return m.group(1).strip(), m.group(2)
    return ch, ''


def parse_address_parts(address: str) -> dict[str, str]:
    keys = ('region', 'district', 'settlement', 'street', 'house')
    out = {k: '' for k in keys}
    if not address or not str(address).strip():
        return out
    chunks = [c.strip() for c in re.split(r'[,;]', str(address)) if c.strip()]
    if len(chunks) == 1:
        out['settlement'] = chunks[0]
        out['region'] = ''
        return out

    leftovers: list[str] = []
    for raw in chunks:
        ch = raw
        extra_house = ''
        street_part, hh = _split_street_house(ch)
        if hh:
            ch = street_part
            extra_house = hh
        kind = _classify_address_chunk(ch)
        if kind and not out[kind]:
            out[kind] = ch
            if extra_house and not out['house']:
                out['house'] = extra_house
        elif kind == 'house' and not out['house']:
            out['house'] = ch
        else:
            leftovers.append(raw)

    for ch in leftovers:
        kind = _classify_address_chunk(ch)
        if kind and not out[kind]:
            out[kind] = ch
            continue
        if re.match(r'^\d+[a-zA-Zа-яА-ЯЁё/-]*$', ch):
            if not out['house']:
                out['house'] = ch
            else:
                out['house'] = f"{out['house']}/{ch}"
            continue
        for key in keys:
            if not out[key]:
                out[key] = ch
                break
    return out


def _split_recommendations(text: str, lines: int = 3) -> list[str]:
    raw = (text or '').strip()
    if not raw:
        return [''] * lines
    parts = [p.strip() for p in re.split(r'[\r\n]+', raw) if p.strip()]
    if len(parts) == 1 and len(parts[0]) > 120:
        words = parts[0].split()
        parts = []
        chunk: list[str] = []
        for w in words:
            chunk.append(w)
            if len(' '.join(chunk)) > 70:
                parts.append(' '.join(chunk))
                chunk = []
        if chunk:
            parts.append(' '.join(chunk))
    while len(parts) < lines:
        parts.append('')
    return parts[:lines]


def build_passport_context(
    *,
    object_row: dict,
    client_row: Optional[dict],
    survey_row: Optional[dict],
    overrides: Optional[dict] = None,
    inline_inputs: Optional[dict] = None,
    inline_computed: Optional[dict] = None,
    inline_conclusion: Optional[str] = None,
) -> dict[str, Any]:
    overrides = overrides or {}
    inputs: dict = {}
    computed: dict = {}
    conclusion = ''

    if survey_row:
        inputs = dict(survey_row.get('inputs') or {})
        computed = dict(survey_row.get('computed') or {})
        conclusion = (survey_row.get('conclusion') or '').strip()
    if inline_inputs:
        inputs.update(inline_inputs)
    if inline_computed:
        computed.update(inline_computed)
    if inline_conclusion is not None:
        conclusion = inline_conclusion.strip()

    address = _pick_str(overrides, inputs, 'address')
    if not address and client_row:
        address = (client_row.get('address') or '').strip()
    if not address:
        address = str(object_row.get('name') or '').strip()

    addr_parts = parse_address_parts(address)
    for key in ('region', 'district', 'settlement', 'street', 'house'):
        v = _pick_str(overrides, inputs, key)
        if v:
            addr_parts[key] = v

    static_m = _num(inputs.get('static_m'))
    dynamic_m = _num(inputs.get('dynamic_m'))
    depth_m = _num(inputs.get('well_depth_m'))
    calc_depth_m = _num(inputs.get('calc_depth_m'))
    measure_sec = _num(inputs.get('measure_seconds'))

    debit_dyn = _num(computed.get('debit_m3h_dynamic'))
    if debit_dyn is None:
        debit_dyn = _debit_from_measure(inputs)

    pump_mark = _pick_str(overrides, inputs, 'pump_mark')
    pipe_text = _pick_str(overrides, inputs, 'pipe')
    casing = _pick_str(overrides, inputs, 'casing_diameter')
    if not casing and pipe_text:
        m = re.search(r'ду\s*(\d+)', pipe_text, re.I)
        if m:
            casing = f'{m.group(1)} мм'
        elif '89' in pipe_text:
            casing = '89 мм'

    material = _pick_str(overrides, inputs, 'pipe_material') or 'сталь'
    filter_iv = _pick_str(overrides, inputs, 'filter_interval')
    sump_iv = _pick_str(overrides, inputs, 'sump_interval')
    executor = _pick_str(overrides, inputs, 'executor') or DEFAULT_EXECUTOR
    rec_lines = _split_recommendations(
        _pick_str(overrides, inputs, 'conclusion') or conclusion, 3
    )

    depth_s = f'{fmt_num_comma(depth_m)}м' if depth_m is not None else ''
    static_s = f'{fmt_num_comma(static_m)}м' if static_m is not None else ''
    dynamic_s = f'{fmt_num_comma(dynamic_m)}м' if dynamic_m is not None else ''
    calc_s = f'{fmt_num_comma(calc_depth_m, 0)}м' if calc_depth_m is not None else ''
    debit_s = f'{fmt_num_comma(debit_dyn, 1)} м3/ч' if debit_dyn is not None else ''
    duration = format_duration_ru(measure_sec)

    pump_display = pump_mark
    if not pump_display:
        p_m3h = _num(inputs.get('pump_m3h'))
        if p_m3h is not None:
            pump_display = f'{fmt_num_comma(p_m3h)} м³/ч'

    return {
        'lines': {
            4: _field_line('Область:', addr_parts['region']),
            5: _field_line('Район:', addr_parts['district']),
            6: _field_line('Нас. Пункт:', addr_parts['settlement']),
            7: _field_line('Улица:', addr_parts['street']),
            8: _field_line('Дом (номер участка):', addr_parts['house']),
            10: _field_line('Диаметр обсадной трубы:', casing),
            11: _field_line('Труба изготовлена из материала:', material),
            12: _field_line('Общая глубина скважины:', depth_s),
            13: _field_line('Интервал установки фильтра:', filter_iv),
            14: _field_line('Интервал установки отстойника:', sump_iv),
            16: _field_line('Статический уровень:', static_s),
            17: _field_line('Динамический уровень:', dynamic_s),
            18: _field_line('Откачка производилась насосам:', pump_display),
            19: _field_line('в течении:', duration),
            20: _field_line('с глубины:', calc_s),
            21: _field_line('с производительностью:', debit_s),
            23: pump_display or '_' * 52,
            25: calc_s or '_' * 52,
            27: rec_lines[0] or '_' * 52,
            28: rec_lines[1] or '_' * 52,
            29: rec_lines[2] or '_' * 52,
            32: executor,
        },
    }


def _set_paragraph_text(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ''


def fill_passport_docx(context: dict[str, Any], dest_path: Path) -> Path:
    if not TEMPLATE_PATH.is_file():
        raise FileNotFoundError(f'Шаблон паспорта не найден: {TEMPLATE_PATH}')
    doc = Document(str(TEMPLATE_PATH))
    paras = doc.paragraphs
    for idx, text in context.get('lines', {}).items():
        if idx < len(paras):
            _set_paragraph_text(paras[idx], text)
    dest_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dest_path))
    return dest_path


def convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> Path:
    docx_path = Path(docx_path)
    pdf_path = Path(pdf_path)
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        from docx2pdf import convert as docx2pdf_convert

        docx2pdf_convert(str(docx_path), str(pdf_path))
        if pdf_path.is_file() and pdf_path.stat().st_size > 0:
            return pdf_path
    except Exception as e:
        logging.debug('docx2pdf: %s', e)
    soffice = shutil.which('soffice') or shutil.which('libreoffice')
    if soffice:
        out_dir = pdf_path.parent
        proc = subprocess.run(
            [
                soffice, '--headless', '--norestore', '--convert-to', 'pdf',
                '--outdir', str(out_dir), str(docx_path),
            ],
            capture_output=True,
            text=True,
            timeout=120,
        )
        if proc.returncode != 0:
            logging.warning('LibreOffice PDF: %s', proc.stderr or proc.stdout)
        generated = out_dir / (docx_path.stem + '.pdf')
        if generated.is_file() and generated != pdf_path:
            generated.replace(pdf_path)
        if pdf_path.is_file() and pdf_path.stat().st_size > 0:
            return pdf_path
    raise PdfConversionError(
        'Не удалось создать PDF. Скачайте DOCX или установите Microsoft Word / LibreOffice.'
    )


def generate_passport_files(
    context: dict[str, Any],
    *,
    want_pdf: bool = True,
    basename: str = 'pasport_skvazhiny',
) -> tuple[Path, Optional[Path]]:
    tmp = Path(tempfile.mkdtemp(prefix='well_passport_'))
    safe = re.sub(r'[^\w\-]+', '_', basename, flags=re.UNICODE)[:80].strip('_') or 'pasport'
    docx_path = tmp / f'{safe}.docx'
    fill_passport_docx(context, docx_path)
    pdf_path = None
    if want_pdf:
        pdf_path = tmp / f'{safe}.pdf'
        try:
            convert_docx_to_pdf(docx_path, pdf_path)
        except PdfConversionError:
            pdf_path = None
    return docx_path, pdf_path


def survey_row_for_passport(row: dict) -> dict:
    def _json_field(val, default):
        if isinstance(val, dict):
            return val
        if not val:
            return default
        try:
            return json.loads(val)
        except Exception:
            return default

    return {
        'measured_at': row.get('measured_at'),
        'conclusion': row.get('conclusion'),
        'inputs': _json_field(row.get('inputs_json'), {}),
        'computed': _json_field(row.get('computed_json'), {}),
    }
