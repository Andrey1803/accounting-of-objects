"""
Автозаполнение акта обследования скважины (DOCX) и экспорт в PDF.
Шаблон: static/templates/well_inspection_act_template.docx
"""
from __future__ import annotations

import logging
import os
import re
import shutil
import subprocess
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

from docx import Document

_BASE_DIR = Path(__file__).resolve().parent
TEMPLATE_PATH = _BASE_DIR / 'static' / 'templates' / 'well_inspection_act_template.docx'

MONTHS_RU_GEN = (
    'января', 'февраля', 'марта', 'апреля', 'мая', 'июня',
    'июля', 'августа', 'сентября', 'октября', 'ноября', 'декабря',
)

DEFAULT_EXECUTOR = os.environ.get(
    'WELL_ACT_EXECUTOR', 'ИП Емельянов Андрей Николаевич'
)
DEFAULT_PHONE = os.environ.get('WELL_ACT_PHONE', '')
DEFAULT_PIPE = os.environ.get(
    'WELL_ACT_PIPE',
    'Труба водоподъёмная ду 89 общей длинной 36м. Способ присоединения фланцевый.',
)
DEFAULT_PUMP_MARK = os.environ.get('WELL_ACT_PUMP_MARK', 'ЭЦВ 8-25-120')
DEFAULT_VIDEO_NOTE = os.environ.get(
    'WELL_ACT_VIDEO_NOTE',
    'В скважине на глубине {depth} м просматриваются илистые отложения, '
    'из-за которых проводить видео инспекцию глубже не представляется возможным.',
)


class PdfConversionError(RuntimeError):
    pass


def _num(v: Any) -> Optional[float]:
    if v is None or v == '':
        return None
    try:
        x = float(str(v).replace(',', '.'))
        return x if x == x else None  # NaN
    except (TypeError, ValueError):
        return None


def fmt_num_comma(val: Optional[float], decimals: int = 1) -> str:
    if val is None:
        return '—'
    s = f'{val:.{decimals}f}'.rstrip('0').rstrip('.')
    return s.replace('.', ',')


def parse_measured_at(value: Any) -> datetime:
    if not value:
        return datetime.now()
    s = str(value).strip().replace('T', ' ')[:19]
    for fmt in ('%Y-%m-%d %H:%M:%S', '%Y-%m-%d %H:%M', '%Y-%m-%d', '%d.%m.%Y'):
        try:
            return datetime.strptime(s[: len(fmt)], fmt)
        except ValueError:
            continue
    try:
        return datetime.fromisoformat(s)
    except ValueError:
        return datetime.now()


def format_date_ru(dt: datetime) -> str:
    return f'{dt.day} {MONTHS_RU_GEN[dt.month - 1]} {dt.year}г.'


def format_duration_ru(seconds: Optional[float]) -> str:
    if seconds is None or seconds <= 0:
        return '1 часа'
    sec = float(seconds)
    if sec >= 3600:
        h = sec / 3600
        hi = int(round(h))
        if hi == 1:
            return '1 часа'
        if 2 <= hi <= 4:
            return f'{hi} часа'
        return f'{hi} часов'
    minutes = max(1, int(round(sec / 60)))
    if minutes == 1:
        return '1 минуты'
    if 2 <= minutes <= 4:
        return f'{minutes} минуты'
    return f'{minutes} минут'


def _debit_from_measure(inputs: dict) -> Optional[float]:
    liters = _num(inputs.get('measure_liters'))
    seconds = _num(inputs.get('measure_seconds'))
    if liters is None or seconds is None or seconds <= 0:
        return None
    return (liters / 1000) / (seconds / 3600)


def _debit_from_pump(inputs: dict) -> Optional[float]:
    stat = _num(inputs.get('static_m'))
    dyn = _num(inputs.get('dynamic_m'))
    calc = _num(inputs.get('calc_depth_m'))
    pump = _num(inputs.get('pump_m3h'))
    if None in (stat, dyn, calc, pump) or abs(dyn - stat) < 1e-9:
        return None
    return pump / (dyn - stat) * (calc - stat)


def build_act_context(
    *,
    object_row: dict,
    client_row: Optional[dict],
    survey_row: Optional[dict],
    overrides: Optional[dict] = None,
    inline_inputs: Optional[dict] = None,
    inline_computed: Optional[dict] = None,
    inline_conclusion: Optional[str] = None,
) -> dict[str, Any]:
    """Собрать текстовые поля акта из объекта, клиента и замера."""
    overrides = overrides or {}
    inputs = {}
    computed = {}
    conclusion = ''
    measured_at = None

    if survey_row:
        inputs = dict(survey_row.get('inputs') or {})
        computed = dict(survey_row.get('computed') or {})
        conclusion = (survey_row.get('conclusion') or '').strip()
        measured_at = survey_row.get('measured_at')
    if inline_inputs:
        inputs.update(inline_inputs)
    if inline_computed:
        computed.update(inline_computed)
    if inline_conclusion is not None:
        conclusion = inline_conclusion.strip()

    dt = parse_measured_at(overrides.get('measured_at') or measured_at)
    customer = (overrides.get('customer') or '').strip()
    if not customer:
        if client_row and client_row.get('name'):
            customer = str(client_row['name']).strip()
        elif object_row.get('client'):
            customer = str(object_row['client']).strip()
    if not customer:
        customer = str(object_row.get('name') or 'Заказчик').strip()

    address = (overrides.get('address') or '').strip()
    if not address and client_row:
        address = (client_row.get('address') or '').strip()
    if not address:
        address = str(object_row.get('name') or '').strip()

    static_m = _num(inputs.get('static_m'))
    dynamic_m = _num(inputs.get('dynamic_m'))
    depth_m = _num(inputs.get('well_depth_m'))
    calc_depth_m = _num(inputs.get('calc_depth_m'))
    measure_sec = _num(inputs.get('measure_seconds'))

    debit_dyn = _num(computed.get('debit_m3h_dynamic'))
    if debit_dyn is None:
        debit_dyn = _debit_from_measure(inputs)

    debit_max = _num(computed.get('debit_m3h'))
    if debit_max is None:
        debit_max = _debit_from_pump(inputs)

    specific = None
    if debit_dyn is not None and static_m is not None and dynamic_m is not None:
        drawdown = dynamic_m - static_m
        if drawdown > 1e-9:
            specific = debit_dyn / drawdown

    pipe = (overrides.get('pipe') or '').strip() or DEFAULT_PIPE
    pump_raw = (overrides.get('pump_mark') or '').strip() or DEFAULT_PUMP_MARK
    pump_line = pump_raw if pump_raw.lower().startswith('установлен') else f'Установлен насос марки {pump_raw}.'

    depth_fmt = fmt_num_comma(depth_m) if depth_m is not None else '—'
    video_default = DEFAULT_VIDEO_NOTE.format(depth=depth_fmt)
    video_note = (overrides.get('video_note') or '').strip() or video_default

    phone = (overrides.get('phone') or '').strip()
    if not phone and client_row:
        phone = (client_row.get('phone') or '').strip()
    if not phone:
        phone = DEFAULT_PHONE
    phone_line = phone
    if phone and not phone.lower().startswith('viber'):
        phone_line = f'Видео материал прилагается и выслан на Viber {phone}'

    executor = (overrides.get('executor') or '').strip() or DEFAULT_EXECUTOR
    conclusion_text = (overrides.get('conclusion') or '').strip() or conclusion
    if conclusion_text and not conclusion_text.lower().startswith('заключение'):
        conclusion_text = f'Заключение: {conclusion_text}'

    dyn_str = fmt_num_comma(dynamic_m)
    debit_dyn_str = fmt_num_comma(debit_dyn, 1)
    duration = format_duration_ru(measure_sec)
    specific_str = fmt_num_comma(specific, 2)
    calc_str = fmt_num_comma(calc_depth_m, 0) if calc_depth_m is not None else '—'
    debit_max_str = fmt_num_comma(debit_max, 0) if debit_max is not None else '—'

    return {
        'date_ru': format_date_ru(dt),
        'customer': customer,
        'address': address,
        'pipe': pipe,
        'pump_line': pump_line,
        'static_line': f'Статический уровень {fmt_num_comma(static_m)}м',
        'depth_line': f'Промеряемая глубина скважины {depth_fmt}м',
        'dynamic_line': (
            f'Динамический уровень {dyn_str} при отборе воды из скважины в течении '
            f'{duration} с производительностью насоса {debit_dyn_str} м3/ч.'
            if dynamic_m is not None and debit_dyn is not None
            else f'Динамический уровень {dyn_str} при отборе воды из скважины.'
        ),
        'specific_line': (
            f'Удельный дебет скважины {specific_str} м/3 на 1 метр понижения зеркала воды.'
            if specific is not None
            else 'Удельный дебет скважины — м/3 на 1 метр понижения зеркала воды.'
        ),
        'max_debit_line': (
            f'Максимальный дебет скважины с глубины {calc_str}м {debit_max_str} м3/ч.'
            if calc_depth_m is not None and debit_max is not None
            else 'Максимальный дебет скважины — м3/ч.'
        ),
        'video_note': video_note,
        'phone_line': phone_line,
        'conclusion': conclusion_text or (
            'Заключение: по результатам обследования рекомендации уточняются.'
        ),
        'signature_line': (
            f'Обследование произвел {executor} ________________________             '
            f'{format_date_ru(dt)}'
        ),
    }


def _set_paragraph_text(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ''


def fill_inspection_act_docx(context: dict[str, Any], dest_path: Path) -> Path:
    if not TEMPLATE_PATH.is_file():
        raise FileNotFoundError(f'Шаблон акта не найден: {TEMPLATE_PATH}')
    doc = Document(str(TEMPLATE_PATH))
    paras = doc.paragraphs
    mapping = {
        1: f"От {context['date_ru']}.",
        2: f"На объекте {context['customer']} по адресу: {context['address']}",
        4: context['pipe'],
        5: context['pump_line'],
        7: context['static_line'],
        8: context['depth_line'],
        9: context['dynamic_line'],
        10: context['specific_line'],
        11: context['max_debit_line'],
        13: context['video_note'],
        14: context['phone_line'],
        17: context['conclusion'],
        24: context['signature_line'],
    }
    for idx, text in mapping.items():
        if idx < len(paras):
            _set_paragraph_text(paras[idx], text)
    dest_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dest_path))
    return dest_path


def convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> Path:
    """DOCX → PDF: docx2pdf (Windows + Word) или LibreOffice."""
    docx_path = Path(docx_path)
    pdf_path = Path(pdf_path)
    pdf_path.parent.mkdir(parents=True, exist_ok=True)

    try:
        from docx2pdf import convert as docx2pdf_convert

        docx2pdf_convert(str(docx_path), str(pdf_path))
        if pdf_path.is_file() and pdf_path.stat().st_size > 0:
            return pdf_path
    except Exception as e:
        logging.debug('docx2pdf: %s', e)

    soffice = shutil.which('soffice') or shutil.which('libreoffice')
    if soffice:
        out_dir = pdf_path.parent
        proc = subprocess.run(
            [
                soffice,
                '--headless',
                '--norestore',
                '--convert-to',
                'pdf',
                '--outdir',
                str(out_dir),
                str(docx_path),
            ],
            capture_output=True,
            text=True,
            timeout=120,
        )
        if proc.returncode != 0:
            logging.warning('LibreOffice PDF: %s', proc.stderr or proc.stdout)
        generated = out_dir / (docx_path.stem + '.pdf')
        if generated.is_file() and generated != pdf_path:
            generated.replace(pdf_path)
        if pdf_path.is_file() and pdf_path.stat().st_size > 0:
            return pdf_path

    raise PdfConversionError(
        'Не удалось создать PDF. Скачайте DOCX или установите Microsoft Word '
        '(Windows) / LibreOffice на сервере.'
    )


def generate_inspection_act_files(
    context: dict[str, Any],
    *,
    want_pdf: bool = True,
    basename: str = 'akt_obledovaniya_skvazhiny',
) -> tuple[Path, Optional[Path]]:
    """Вернуть (docx_path, pdf_path|None) во временной папке."""
    tmp = Path(tempfile.mkdtemp(prefix='well_act_'))
    safe = re.sub(r'[^\w\-]+', '_', basename, flags=re.UNICODE)[:80].strip('_') or 'akt'
    docx_path = tmp / f'{safe}.docx'
    fill_inspection_act_docx(context, docx_path)
    pdf_path = None
    if want_pdf:
        pdf_path = tmp / f'{safe}.pdf'
        try:
            convert_docx_to_pdf(docx_path, pdf_path)
        except PdfConversionError:
            pdf_path = None
    return docx_path, pdf_path


def survey_row_for_act(row: dict) -> dict:
    """Нормализовать строку БД (inputs_json / computed_json) для build_act_context."""
    import json

    def _json_field(val, default):
        if isinstance(val, dict):
            return val
        if not val:
            return default
        try:
            return json.loads(val)
        except Exception:
            return default

    return {
        'measured_at': row.get('measured_at'),
        'conclusion': row.get('conclusion'),
        'inputs': _json_field(row.get('inputs_json'), {}),
        'computed': _json_field(row.get('computed_json'), {}),
    }
